import streamlit as st
import PyPDF2
import docx
import os
import io
import base64
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import anthropic
from dotenv import load_dotenv

# Laad API sleutel
load_dotenv()
anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")

# Configuratie
st.set_page_config(page_title="Document Summarizer", layout="wide")

# Styling
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stButton button {
        width: 100%;
        background-color: #4CAF50;
        color: white;
    }
    .stTitle {
        font-size: 2.5rem;
        font-weight: bold;
        color: #2E2E2E;
    }
</style>
""", unsafe_allow_html=True)

# Titel
st.title("Document Summarizer")
st.write("Upload je documenten, kies je voorkeuren, en ontvang een gepersonaliseerde samenvatting.")

# Status van API-sleutel tonen
if not anthropic_api_key:
    st.error("⚠️ Geen API-sleutel gevonden. Voer de sleutel in om deze app te gebruiken.")
    api_key_input = st.text_input("Voer je Anthropic API-sleutel in:", type="password")
    if api_key_input:
        anthropic_api_key = api_key_input
        st.success("API-sleutel ingesteld!")

# Functies voor document verwerking
def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:  # Controleer of er tekst is geëxtraheerd
                text += page_text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(docx_file):
    """Extract text from a Word document."""
    try:
        doc = docx.Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            if para.text:  # Controleer of er tekst is
                text += para.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def process_documents(uploaded_files):
    """Process multiple documents and combine their text."""
    combined_text = ""
    for uploaded_file in uploaded_files:
        # Maak een tijdelijke kopie van het bestand
        bytes_data = uploaded_file.getvalue()
        temp_file = io.BytesIO(bytes_data)
        temp_file.name = uploaded_file.name
        
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension == ".pdf":
            text = extract_text_from_pdf(temp_file)
        elif file_extension in [".docx", ".doc"]:
            text = extract_text_from_docx(temp_file)
        else:
            st.warning(f"Bestandstype {file_extension} wordt niet ondersteund.")
            continue
            
        combined_text += f"\n\n--- Document: {uploaded_file.name} ---\n\n{text}"
    
    return combined_text

def create_pdf(text, filename="samenvatting.pdf"):
    """Create a PDF file from text."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Configure text properties
    c.setFont("Helvetica", 12)
    y_position = height - 50
    lines = text.split('\n')
    
    for line in lines:
        # Check if we need a new page
        if y_position < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y_position = height - 50
        
        # Split long lines
        words = line.split()
        current_line = ""
        for word in words:
            if c.stringWidth(current_line + word, "Helvetica", 12) < width - 100:
                current_line += word + " "
            else:
                c.drawString(50, y_position, current_line)
                y_position -= 15
                current_line = word + " "
        
        if current_line:
            c.drawString(50, y_position, current_line)
            y_position -= 15
    
    c.save()
    buffer.seek(0)
    return buffer

def create_docx(text, filename="samenvatting.docx"):
    """Create a Word document from text."""
    doc = Document()
    doc.add_heading("Samenvatting", 0)
    
    paragraphs = text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_download_link(buffer, filename, format_type):
    """Generate a download link for the file."""
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode()
    
    if format_type == "pdf":
        mime_type = "application/pdf"
    else:  # docx
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}">Download {filename}</a>'
    return href

def summarize_text(text, summary_length, style, formality, purpose, industry):
    """Generate a summary using the Claude API."""
    try:
        # Maak een nieuwe client aan met de API-sleutel
        client = anthropic.Anthropic(api_key=anthropic_api_key)
        
        # Construct the prompt
        prompt = f"""
        Als een expert in het maken van samenvattingen, verwerk de volgende tekst:
        
        [BEGIN TEKST]
        {text}
        [EINDE TEKST]
        
        Maak een samenvatting met de volgende kenmerken:
        - Lengte: {summary_length}
        - Stijl: {style}
        - Formaliteit: {formality}
        - Doel: {purpose}
        - Branche: {industry}
        
        Zorg ervoor dat je de belangrijkste punten benadrukt en dubbelingen vermijdt. Geef als resultaat een gestructureerde, goed leesbare samenvatting.
        """

        # Call the Claude API
        message = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=4000,
            temperature=0.3,
            system="Je bent een expert in het maken van samenvattingen. Je maakt beknopte en informatieve samenvattingen van documenten.",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        return message.content[0].text

    except Exception as e:
        st.error(f"❌ Error generating summary: {e}")
        return "Er is een fout opgetreden bij het genereren van de samenvatting. Probeer het opnieuw."

# Structuur van de app
with st.sidebar:
    st.header("Documentinvoer")
    uploaded_files = st.file_uploader("Upload PDF of Word bestanden", 
                                     type=["pdf", "docx", "doc"], 
                                     accept_multiple_files=True)
    
    st.header("Voorkeuren")
    
    summary_length = st.selectbox(
        "Lengte samenvatting",
        ["Kort (10-15% van origineel)", "Medium (20-25% van origineel)", "Lang (30-40% van origineel)", "Uitgebreid (50% van origineel)"]
    )
    
    style = st.selectbox(
        "Stijl",
        ["Zakelijk", "Academisch", "Informeel", "Journalistiek"]
    )
    
    formality = st.selectbox(
        "Formaliteit",
        ["Zeer formeel", "Formeel", "Semi-formeel", "Informeel"]
    )
    
    purpose = st.selectbox(
        "Doel",
        ["Informatief", "Besluitvorming", "Presentatie", "Training"]
    )
    
    industry = st.selectbox(
        "Branche",
        ["Algemeen", "Technologie", "Financiën", "Gezondheidszorg"]
    )
    
    output_format = st.radio(
        "Output formaat",
        ["PDF", "Word"]
    )

# Hoofdsectie
if uploaded_files:
    st.header("Verwerking")
    
    # Verwerk de documenten
    with st.spinner("Documenten verwerken..."):
        combined_text = process_documents(uploaded_files)
    
    st.success(f"Succesvol {len(uploaded_files)} document(en) verwerkt.")
    
    # Toon een voorbeeld van de gecombineerde tekst
    with st.expander("Bekijk ruwe tekst"):
        st.text_area("Gecombineerde tekst", combined_text, height=200)
    
    # Genereer samenvatting
    if st.button("Genereer samenvatting"):
        if not anthropic_api_key:
            st.error("⚠️ API-sleutel ontbreekt. Voer deze in om samenvattingen te genereren.")
        else:
            with st.spinner("Samenvatting genereren..."):
                summary = summarize_text(
                    combined_text, 
                    summary_length, 
                    style, 
                    formality, 
                    purpose, 
                    industry
                )
            
            st.header("Samenvatting")
            st.write(summary)
            
            # Download opties
            st.header("Download opties")
            
            if output_format == "PDF":
                pdf_buffer = create_pdf(summary)
                st.markdown(get_download_link(pdf_buffer, "samenvatting.pdf", "pdf"), unsafe_allow_html=True)
            else:  # Word
                docx_buffer = create_docx(summary)
                st.markdown(get_download_link(docx_buffer, "samenvatting.docx", "docx"), unsafe_allow_html=True)
else:
    st.info("Upload één of meerdere documenten om te beginnen.")

# Voeg wat extra informatie toe
st.sidebar.markdown("---")
st.sidebar.markdown("### Over deze app")
st.sidebar.write("Deze app gebruikt AI om documenten samen te vatten en te optimaliseren volgens jouw voorkeuren.")
